"""Generate the final DOCX report for submission."""
# ruff: noqa: E402

from __future__ import annotations

import json
import os
from datetime import date
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", str((Path.cwd() / ".cache" / "matplotlib").resolve()))
(Path.cwd() / ".cache" / "matplotlib").mkdir(parents=True, exist_ok=True)

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from heart_disease_mlops.config import FIGURES_DIR, METRICS_PATH, REPORTS_DIR

TITLE = "MLOps Assignment I: Heart Disease Risk Prediction"
SUBTITLE = "End-to-End ML Development, CI/CD, Containerization, Deployment, and Monitoring"
OUTPUT_PATH = REPORTS_DIR / "MLOps_Assignment_Report.docx"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_run_font(run, size: float = 11, bold: bool | None = None, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "BITS MLOps Assignment I | Heart Disease UCI"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.text = "Repository: https://github.com/tosenthu/bits-mlops"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color="666666")


def add_title_page(doc: Document) -> None:
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(TITLE)
    set_run_font(title_run, size=23, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(SUBTITLE)
    set_run_font(subtitle_run, size=13, color="333333")

    meta = doc.add_table(rows=5, cols=2)
    meta.autofit = False
    widths = [Inches(1.8), Inches(4.2)]
    rows = [
        ("Course", "MLOps (S2-25_AMLCSZG523)"),
        ("Dataset", "Heart Disease UCI - Cleveland processed dataset"),
        ("Problem", "Binary classification for heart disease risk prediction"),
        ("Repository", "https://github.com/tosenthu/bits-mlops"),
        ("Date", date.today().isoformat()),
    ]
    for row_index, (label, value) in enumerate(rows):
        cells = meta.rows[row_index].cells
        cells[0].width = widths[0]
        cells[1].width = widths[1]
        shade_cell(cells[0], "F2F4F7")
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run(
        "This report documents a reproducible machine learning system with automated "
        "data preparation, experiment tracking, model packaging, API serving, CI/CD, "
        "containerization, Kubernetes deployment artifacts, and monitoring hooks."
    )
    set_run_font(note_run, size=11)
    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for column_index, header in enumerate(headers):
        cell = table.rows[0].cells[column_index]
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for column_index, value in enumerate(row):
            set_cell_text(cells[column_index], value)


def read_metrics() -> dict:
    if METRICS_PATH.exists():
        return json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    return {
        "models": [],
        "best_model": "Run `python -m heart_disease_mlops.train --fast` to generate metrics.",
    }


def create_architecture_diagram(output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig, axis = plt.subplots(figsize=(10, 4.8))
    axis.axis("off")
    boxes = [
        ("UCI Dataset", 0.05, 0.62),
        ("EDA + Cleaning", 0.23, 0.62),
        ("Feature Pipeline", 0.41, 0.62),
        ("Model Training\nMLflow", 0.59, 0.62),
        ("FastAPI + Docker", 0.77, 0.62),
        ("Kubernetes\nService", 0.77, 0.25),
        ("Prometheus\nMetrics", 0.59, 0.25),
        ("GitHub Actions\nCI/CD", 0.41, 0.25),
    ]
    for label, x_coord, y_coord in boxes:
        axis.text(
            x_coord,
            y_coord,
            label,
            ha="center",
            va="center",
            fontsize=10,
            bbox={"boxstyle": "round,pad=0.35", "facecolor": "#F2F4F7", "edgecolor": "#5B6B7C"},
        )
    arrows = [
        ((0.12, 0.62), (0.18, 0.62)),
        ((0.30, 0.62), (0.36, 0.62)),
        ((0.48, 0.62), (0.54, 0.62)),
        ((0.66, 0.62), (0.72, 0.62)),
        ((0.77, 0.55), (0.77, 0.33)),
        ((0.72, 0.25), (0.66, 0.25)),
        ((0.54, 0.25), (0.48, 0.25)),
    ]
    for start, end in arrows:
        axis.annotate("", xy=end, xytext=start, arrowprops={"arrowstyle": "->", "lw": 1.5})
    plt.tight_layout()
    plt.savefig(output_path, dpi=180)
    plt.close()
    return output_path


def add_picture_if_exists(doc: Document, path: Path, width: float = 6.0) -> None:
    if path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(width))
    else:
        doc.add_paragraph(f"Figure pending: {path.name}")


def add_report_body(doc: Document) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The project implements a complete MLOps workflow for predicting heart disease risk "
        "from the UCI Cleveland dataset. The solution includes dataset acquisition, EDA, "
        "preprocessing, model training, experiment tracking, model persistence, API serving, "
        "Docker packaging, Kubernetes deployment manifests, CI/CD, logging, and metrics."
    )
    add_bullets(
        doc,
        [
            "Problem type: supervised binary classification.",
            "Primary models: Logistic Regression and Random Forest.",
            "Serving pattern: FastAPI application with /predict, /health, and /metrics endpoints.",
            "Deployment pattern: Docker container deployed through Kubernetes manifests.",
        ],
    )

    doc.add_heading("2. Dataset and Data Acquisition", level=1)
    doc.add_paragraph(
        "The dataset is the processed Cleveland Heart Disease dataset from the UCI Machine "
        "Learning Repository. The repository includes scripts/download_data.py, which downloads "
        "the source CSV, applies column names, converts '?' values to missing values, converts "
        "columns to numeric types, and binarizes the target."
    )
    add_table(
        doc,
        ["Field group", "Columns", "Treatment"],
        [
            ["Demographic", "age, sex", "Age is scaled; sex is one-hot encoded."],
            ["Clinical", "trestbps, chol, fbs, thalach", "Numeric values are imputed and scaled."],
            [
                "ECG and symptoms",
                "cp, restecg, exang, oldpeak, slope",
                "Mixed numeric and categorical preprocessing.",
            ],
            ["Diagnostic", "ca, thal", "Missing values are imputed inside the pipeline."],
            ["Target", "target", "Converted to 0 for absence and 1 for presence of disease."],
        ],
    )
    doc.add_page_break()

    doc.add_heading("3. Exploratory Data Analysis", level=1)
    doc.add_paragraph(
        "EDA focuses on target balance, feature distributions, and numeric correlations. The "
        "figures are generated by scripts/run_eda.py and saved under reports/figures."
    )
    add_picture_if_exists(doc, FIGURES_DIR / "class_balance.png", width=4.8)
    add_picture_if_exists(doc, FIGURES_DIR / "feature_histograms.png", width=6.2)
    add_picture_if_exists(doc, FIGURES_DIR / "correlation_heatmap.png", width=6.1)
    doc.add_page_break()

    doc.add_heading("4. Feature Engineering and Model Development", level=1)
    doc.add_paragraph(
        "The repository uses a scikit-learn Pipeline and ColumnTransformer so the exact same "
        "preprocessing steps are used in training and inference. Numeric fields are imputed "
        "with the median and standardized. Categorical fields are imputed with the most frequent "
        "value and one-hot encoded."
    )
    add_table(
        doc,
        ["Model", "Tuning parameters", "Selection reason"],
        [
            [
                "Logistic Regression",
                "C over 0.1, 1.0, 10.0",
                "Interpretable linear baseline with class weighting.",
            ],
            [
                "Random Forest",
                "n_estimators, max_depth, min_samples_leaf",
                "Nonlinear ensemble baseline with robust tabular performance.",
            ],
        ],
    )
    metrics = read_metrics()
    model_rows = []
    for model in metrics.get("models", []):
        test_metrics = model.get("test_metrics", {})
        model_rows.append(
            [
                model.get("model_name", ""),
                f"{test_metrics.get('accuracy', 0):.3f}",
                f"{test_metrics.get('precision', 0):.3f}",
                f"{test_metrics.get('recall', 0):.3f}",
                f"{test_metrics.get('roc_auc', 0):.3f}",
            ]
        )
    if model_rows:
        add_table(doc, ["Model", "Accuracy", "Precision", "Recall", "ROC-AUC"], model_rows)
        add_picture_if_exists(doc, FIGURES_DIR / "model_comparison.png", width=5.8)
    doc.add_paragraph(f"Selected model: {metrics.get('best_model')}")
    doc.add_page_break()

    doc.add_heading("5. Experiment Tracking", level=1)
    doc.add_paragraph(
        "MLflow is integrated in the training entry point. Each candidate model run logs model "
        "name, hyperparameters, cross-validation ROC-AUC, held-out test metrics, the sklearn "
        "model artifact, metrics JSON, model metadata, and model comparison plot."
    )
    add_numbered(
        doc,
        [
            "Run `python -m heart_disease_mlops.train --fast` for a smoke training run.",
            "Run `mlflow ui --backend-store-uri mlruns` to inspect local experiments.",
            "Capture screenshots of run metrics and artifacts into reports/screenshots.",
        ],
    )

    doc.add_heading("6. Packaging and Reproducibility", level=1)
    doc.add_paragraph(
        "The final model is saved as a joblib artifact containing the full preprocessing and "
        "classification pipeline. requirements.txt and pyproject.toml define the runtime and "
        "test environment. The sample_input/patient_high_risk.json file can be used for a "
        "repeatable API check."
    )
    add_bullets(
        doc,
        [
            "models/heart_disease_model.joblib: trained pipeline.",
            "models/model_metadata.json: feature order, selected model, parameters, and metrics.",
            "data/processed/heart_disease_clean.csv: cleaned dataset generated from source.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("7. CI/CD and Automated Testing", level=1)
    doc.add_paragraph(
        "The GitHub Actions workflow in .github/workflows/ci.yml installs dependencies, runs "
        "Ruff linting, executes pytest, performs a fast model training smoke test, and uploads "
        "model, metrics, and report artifacts."
    )
    add_table(
        doc,
        ["Stage", "Purpose", "Failure behavior"],
        [
            [
                "Lint",
                "Runs ruff against source, scripts, and tests.",
                "Fails on style or code-quality issues.",
            ],
            [
                "Unit tests",
                "Validates data cleaning, feature preprocessing, and API behavior.",
                "Fails on regression or schema mismatch.",
            ],
            [
                "Training smoke",
                "Runs a reduced hyperparameter grid.",
                "Fails if data, MLflow, or model persistence breaks.",
            ],
            [
                "Artifacts",
                "Uploads models, metrics, figures, and report outputs.",
                "Provides audit trail for each workflow run.",
            ],
        ],
    )

    doc.add_heading("8. API, Containerization, and Local Serving", level=1)
    doc.add_paragraph(
        "The serving layer is implemented with FastAPI. Dockerfile builds a lightweight image, "
        "trains or uses the packaged model, exposes port 8000, and serves the API through uvicorn."
    )
    add_numbered(
        doc,
        [
            "Build: `docker build -t bits-mlops-heart .`",
            "Run: `docker run --rm -p 8000:8000 bits-mlops-heart`",
            "Predict: `curl -X POST http://localhost:8000/predict "
            "-H 'Content-Type: application/json' "
            "-d @sample_input/patient_high_risk.json`",
        ],
    )
    doc.add_page_break()

    doc.add_heading("9. Production Deployment and Monitoring", level=1)
    doc.add_paragraph(
        "The deployment/k8s folder contains Kubernetes Deployment, Service, and optional "
        "Ingress manifests. The API exposes /metrics using prometheus-client. Request counts "
        "and latency histograms can be scraped by Prometheus and visualized in Grafana."
    )
    add_picture_if_exists(
        doc,
        create_architecture_diagram(FIGURES_DIR / "architecture.png"),
        width=6.1,
    )
    add_numbered(
        doc,
        [
            "Build and tag the Docker image.",
            "Load the image into Minikube or push it to a registry.",
            "Apply deployment/k8s/*.yaml.",
            "Use kubectl port-forward or the configured service URL to verify "
            "/health and /predict.",
            "Add Prometheus scrape configuration for /metrics.",
        ],
    )

    doc.add_heading("10. Submission Evidence Checklist", level=1)
    add_table(
        doc,
        ["Evidence item", "Repository location", "Status"],
        [
            ["Code repository", "https://github.com/tosenthu/bits-mlops", "Included"],
            ["Dataset scripts", "scripts/download_data.py", "Included"],
            ["EDA figures", "reports/figures", "Generated by script"],
            ["Unit tests", "tests/", "Included"],
            ["CI/CD workflow", ".github/workflows/ci.yml", "Included"],
            ["Dockerfile", "Dockerfile", "Included"],
            ["Kubernetes manifests", "deployment/k8s", "Included"],
            ["Screenshots", "reports/screenshots", "Add after running workflow/deployment"],
            ["Demo video", "External short recording", "Record final run-through"],
        ],
    )

    doc.add_heading("Appendix A. Commands", level=1)
    commands = [
        "python -m pip install -r requirements.txt",
        "python scripts/download_data.py",
        "python scripts/run_eda.py",
        "python -m heart_disease_mlops.train --fast",
        "pytest -q",
        "uvicorn heart_disease_mlops.api:app --host 0.0.0.0 --port 8000",
        "docker build -t bits-mlops-heart .",
        "kubectl apply -f deployment/k8s/",
    ]
    for command in commands:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(command)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def generate_report(output_path: Path = OUTPUT_PATH) -> Path:
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_page(doc)
    add_report_body(doc)

    last_section = doc.add_section(WD_SECTION.CONTINUOUS)
    last_section.top_margin = Inches(1)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main() -> None:
    output_path = generate_report()
    print(output_path)


if __name__ == "__main__":
    main()
